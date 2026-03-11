from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Loo uus dokument
doc = Document()

# Pealkiri
title = doc.add_heading('Poe müüja programmi dokumentatsioon', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Autor
author = doc.add_paragraph()
author.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = author.add_run('Autor: Sinu Nimi')
run.font.size = Pt(12)
run.font.italic = True

doc.add_paragraph()

# Sissejuhatus
doc.add_heading('1. Sissejuhatus', level=1)
doc.add_paragraph(
    'See dokument kirjeldab Pythoni programmi, mis kasutab pygame teeki '
    'poemüüja kuvamiseks graafilises aknas. '
    'Programm loob 640×480 piksli suuruse akna, kus kuvatakse poe taust, '
    'müüja ning jutumull tekstiga.'
)

# Nõuded
doc.add_heading('2. Süsteemi nõuded', level=1)
doc.add_paragraph('Programmi käivitamiseks on vajalikud järgmised komponendid:')
nouded = doc.add_paragraph(style='List Bullet')
nouded.add_run('Python 3.x').bold = True
nouded = doc.add_paragraph(style='List Bullet')
nouded.add_run('pygame teek').bold = True
doc.add_paragraph(
    'Pygame saab installida käsuga: pip install pygame'
)

# Kasutatavad failid
doc.add_heading('3. Kasutatavad failid', level=1)
doc.add_paragraph('Programm kasutab järgmisi pildifaile:')
tbl = doc.add_table(rows=4, cols=3)
tbl.style = 'Light Grid Accent 1'
hdr_cells = tbl.rows[0].cells
hdr_cells[0].text = 'Failinimi'
hdr_cells[1].text = 'Kirjeldus'
hdr_cells[2].text = 'Kasutus'

row = tbl.rows[1].cells
row[0].text = 'bg_shop.jpg'
row[1].text = 'Poe taustapilt (640×480)'
row[2].text = 'Taust'

row = tbl.rows[2].cells
row[0].text = 'seller.png'
row[1].text = 'Müüja pilt'
row[2].text = 'Müüja tegelane'

row = tbl.rows[3].cells
row[0].text = 'chat.png'
row[1].text = 'Jutumulli pilt'
row[2].text = 'Tekstikast'

# Programmi struktuur
doc.add_heading('4. Programmi struktuur', level=1)

doc.add_heading('4.1. Initsialiseerimine', level=2)
doc.add_paragraph(
    'Programmi alguses initsialiseeritakse pygame ja luuakse aken '
    'mõõtmetega 640×480 pikslit. Aknale määratakse pealkirja "Ülesanne 2".'
)

doc.add_heading('4.2. Piltide laadimine', level=2)
doc.add_paragraph(
    'Programm laeb kolm pildifaili: bg_shop.jpg (taust), seller.png (müüja) '
    'ja chat.png (jutumull). Müüja pilti vähendatakse, et see mahuks aknasse.'
)

doc.add_heading('4.3. Värvide määramine', level=2)
doc.add_paragraph('Programmis kasutatakse RGB värvisüsteemi järgmise värviga:')
tbl2 = doc.add_table(rows=2, cols=3)
tbl2.style = 'Light Grid Accent 1'
hdr_cells = tbl2.rows[0].cells
hdr_cells[0].text = 'Värv'
hdr_cells[1].text = 'RGB väärtus'
hdr_cells[2].text = 'Kasutus'

row = tbl2.rows[1].cells
row[0].text = 'Valge'
row[1].text = '(255, 255, 255)'
row[2].text = 'Tekst jutumullis'

doc.add_heading('4.4. Elementide kuvamine', level=2)
doc.add_paragraph(
    'Programm kuvab järgmised elemendid õiges järjekorras:'
)
elemendid = [
    ('Taust', 'bg_shop.jpg kuvatakse täisekraanina (0, 0)'),
    ('Müüja', 'seller.png kuvatakse vasakus alumises nurgas'),
    ('Jutumull', 'chat.png kuvatakse müüja kohal'),
    ('Tekst', 'Valge tekst "Tere, olen Sinu Nimi" jutumulli keskele')
]
for nimi, kirjeldus in elemendid:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(nimi + ': ').bold = True
    p.add_run(kirjeldus)

# Põhisilmus
doc.add_heading('5. Põhisilmus (Main Loop)', level=1)
doc.add_paragraph(
    'Pärast graafika joonistamist käivitub põhisilmus, mis hoiab akna avatuna. '
    'Silmus töötab seni, kuni kasutaja sulgeb akna (pygame.QUIT sündmus). '
    'See tagab, et programm ei lõpetaks kohe tööd, vaid ootaks kasutaja tegevust.'
)

# Kasutatud funktsioonid
doc.add_heading('6. Kasutatud pygame funktsioonid', level=1)
funktsioonid = [
    ('pygame.init()', 'Initsialiseerib pygame teegi'),
    ('pygame.display.set_mode()', 'Loob graafilise akna'),
    ('pygame.display.set_caption()', 'Määrab akna pealkirja'),
    ('pygame.image.load()', 'Laeb pildifaili'),
    ('pygame.transform.scale()', 'Muudab pildi suurust'),
    ('screen.blit()', 'Kuvab pildi ekraanil'),
    ('pygame.font.SysFont()', 'Loob fondi teksti jaoks'),
    ('font.render()', 'Renderdab teksti pildiks'),
    ('pygame.display.flip()', 'Uuendab ekraani sisu'),
    ('pygame.event.get()', 'Võtab sündmused järjekorrast'),
    ('pygame.quit()', 'Lõpetab pygame töö')
]
for funktsioon, kirjeldus in funktsioonid:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(funktsioon).bold = True
    p.add_run(' - ' + kirjeldus)

# Järeldus
doc.add_heading('7. Järeldus', level=1)
doc.add_paragraph(
    'Antud programm demonstreerib pygame teegi võimalusi piltide ja teksti '
    'kuvamiseks graafilises aknas. Programm kasutab erinevaid pildifaile '
    '(taust, tegelane, jutumull) ning lisab neile teksti.'
)

doc.add_paragraph(
    'Edasised täiendused võiksid hõlmata animatsioone, kasutaja interaktsiooni '
    '(näiteks nupuvajutused) või heliefektide lisamist.'
)

# Salvesta dokument
doc.save('Poemüüja_dokumentatsioon.docx')
print('Dokument "Poemüüja_dokumentatsioon.docx" loodud edukalt!')
