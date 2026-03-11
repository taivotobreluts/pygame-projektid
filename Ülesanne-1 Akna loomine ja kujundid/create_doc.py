from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Loo uus dokument
doc = Document()

# Pealkiri
title = doc.add_heading('TARKVARAARENDUSE PROJEKTI DOKUMENTATSIOON', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Autor ja ülesanne info
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('\nÜlesanne 1: Akna loomine ja kujundid\n')
run.font.size = Pt(14)
run.font.bold = True
run = info.add_run('Autor: Taivo Tobreluts\n')
run.font.size = Pt(12)
run.font.italic = True

doc.add_paragraph()

# 1. Praktilise töö sooritamise käik
doc.add_heading('1. Praktilise töö sooritamise käik', level=1)

doc.add_heading('Eesmärk:', level=2)
doc.add_paragraph(
    'Luua PyGame graafiline aken, kus kuvatakse liiklustulede (foori) simulatsioon '
    'kolme erineva värviga tulega: punane, kollane ja roheline. '
    'Õppida pygame teegi põhifunktsioone ja graafiliste kujundite joonistamist.'
)

doc.add_heading('Lühikirjeldus:', level=2)
doc.add_paragraph(
    'Programm loob 300×300 piksli suuruse akna musta taustaga. '
    'Foori raam joonistatakse halli ristkülikuna (100×260 px, joonelaius 2 px), '
    'mille sisse paigutatakse kolm ringi raadiusega 38 px: punane (ülemine), kollane (keskmine) ja roheline (alumine). '
    'Programmi lõpus käivitub while-loop, mis hoiab akna avatuna '
    'kuni kasutaja sulgeb selle (pygame.QUIT sündmus).'
)

doc.add_heading('GitHub link:', level=2)
doc.add_paragraph('(Lisa siia oma GitHub repository link)')

doc.add_heading('Töötasid üksi või meeskonnas?', level=2)
doc.add_paragraph('Töötasin üksi.')

doc.add_heading('Kas said abi või abistasid kaasõppijat?', level=2)
doc.add_paragraph(
    'Antud ülesande raames töötasin iseseisvalt. '
    'Kasutasin õppematerjale ja pygame dokumentatsiooni.'
)

doc.add_page_break()

# 2. TI kasutamine
doc.add_heading('2. TI (tehisintellekti) kasutamine', level=1)

doc.add_heading('Kas kasutasid töö käigus tehisintellekti abi?', level=2)
doc.add_paragraph(
    'Jah, kasutasin AI assistenti (Kimi Code CLI) ülesande mõistmiseks ja '
    'dokumentatsiooni vormistamiseks.'
)

doc.add_heading('Millist TI-d kasutasid?', level=2)
doc.add_paragraph('Kimi Code CLI (Kimi AI)')

doc.add_heading('Millised olid esitatud küsimused või käsud (promptid)?', level=2)
doc.add_paragraph('Näited küsimustest:')
prompts = [
    'Kuidas luua PyGame akent ja joonistada sinna kujundeid?',
    'Kuidas tsentreerida foori raam aknas?',
    'Kuidas joonistada ringe pygame.draw.circle() funktsiooniga?',
    'Millised on pygame põhivärvide RGB väärtused?'
]
for prompt in prompts:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(prompt).italic = True

doc.add_heading('Kuidas hindad saadud abi usaldusväärsust?', level=2)
doc.add_paragraph(
    'Saadud abi oli usaldusväärne. AI selgitas pygame funktsioonide kasutamist õigesti. '
    'Kontrollisin kõik saadud koodinäited pygame ametliku dokumentatsiooni põhjal. '
    'Koordinaatide arvutamisel ja värviväärtuste määramisel kasutasin enda teadmisi '
    'kontrollimaks, kas tulemus on loogiline.'
)

doc.add_page_break()

# 3. Esinenud probleemid ja lahendused
doc.add_heading('3. Esinenud probleemid ja lahendused', level=1)

doc.add_heading('Probleem 1: Foori raami tsentreerimine aknas', level=2)
doc.add_paragraph(
    'Probleem: Kuidas arvutada foori raami X-koordinaat, et see oleks akna keskel?\n'
    'Lahendus: Kasutasin valemit: raami_x = (akna_laius - raami_laius) // 2. '
    'See tagab, et raam asub täpselt keskel, sõltumata akna suurusest.'
)

doc.add_heading('Probleem 2: Ringide paigutamine vertikaalselt', level=2)
doc.add_paragraph(
    'Probleem: Kuidas paigutada kolm ringi üksteise alla võrdsete vahedega?\n'
    'Lahendus: Arvutasin iga ringi Y-koordinaadi raami ülemisest servast lähtudes: '
    'esimene ring 45 px, teine 130 px, kolmas 215 px kaugusel raami ülaservast. '
    'X-koordinaat on alati akna keskpunkt (suurus // 2).'
)

doc.add_heading('Probleem 3: Pygame ühilduvus Python 3.14-ga', level=2)
doc.add_paragraph(
    'Probleem: Pygame teek ei toeta 10.03.26 seisuga Python 3.14 versiooni '
    '(ja paistab, et selle edasine arendus on lõppenud).\n'
    'Lahendus: Kasutan uuemat PyGame Community Edition repot (pygame-ce), '
    'mida saab installida käsuga: pip install pygame-ce. '
    'Repo asub aadressil: https://github.com/pygame-community/pygame-ce'
)

doc.add_heading('Probleem 4: Akna lahtipüsimine', level=2)
doc.add_paragraph(
    'Probleem: Programm lõpetas kohe pärast joonistamist töö.\n'
    'Lahendus: Lisasin pygame põhisilmuse (main loop), mis töötab seni, '
    'kuni kasutaja sulgeb akna. Kasutasin while-tsüklit ja pygame.event.get() '
    'sündmuste jälgimiseks.'
)

# Lisa: Kasutatud värvide tabel
doc.add_page_break()
doc.add_heading('Lisa: Kasutatud värvid', level=1)
doc.add_paragraph('Programmis defineeritakse värvid RGB süsteemis:')

tbl = doc.add_table(rows=6, cols=3)
tbl.style = 'Light Grid Accent 1'
hdr_cells = tbl.rows[0].cells
hdr_cells[0].text = 'Värv'
hdr_cells[1].text = 'RGB väärtus'
hdr_cells[2].text = 'Kasutus'

row = tbl.rows[1].cells
row[0].text = 'Must'
row[1].text = '(0, 0, 0)'
row[2].text = 'Taust'

row = tbl.rows[2].cells
row[0].text = 'Hall'
row[1].text = '(128, 128, 128)'
row[2].text = 'Foori raam'

row = tbl.rows[3].cells
row[0].text = 'Punane'
row[1].text = '(255, 0, 0)'
row[2].text = 'Ülemine tuli'

row = tbl.rows[4].cells
row[0].text = 'Kollane'
row[1].text = '(255, 255, 0)'
row[2].text = 'Keskmine tuli'

row = tbl.rows[5].cells
row[0].text = 'Roheline'
row[1].text = '(0, 255, 0)'
row[2].text = 'Alumine tuli'

# Lisa: Kasutatud pygame funktsioonid
doc.add_heading('Lisa: Kasutatud pygame funktsioonid', level=1)
funktsioonid = [
    ('pygame.init()', 'Initsialiseerib pygame teegi'),
    ('pygame.display.set_mode()', 'Loob graafilise akna'),
    ('pygame.display.set_caption()', 'Määrab akna pealkirja'),
    ('pygame.draw.rect()', 'Joonistab ristküliku'),
    ('pygame.draw.circle()', 'Joonistab ringi'),
    ('pygame.display.flip()', 'Uuendab ekraani sisu'),
    ('pygame.event.get()', 'Võtab sündmused järjekorrast'),
    ('pygame.quit()', 'Lõpetab pygame töö')
]
for funktsioon, kirjeldus in funktsioonid:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(funktsioon + '\t\t- ').bold = True
    p.add_run(kirjeldus)

# Salvesta dokument
doc.save('Foori_dokumentatsioon.docx')
print('Dokument "Foori_dokumentatsioon.docx" loodud edukalt!')
